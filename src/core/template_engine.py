"""Template engine and batch exporter.

Heavy third-party dependencies (Jinja2, WeasyPrint, python-docx, qrcode)
are nice-to-have for the GUI and export functionality. To make lightweight
unit testing (for example `tests/test_boss_fight.py`) possible without
installing all optional dependencies, we import the heavy modules lazily
and fall back to placeholders if they're missing. BatchExporter only uses
the Database API and does not require those heavy imports.
"""

try:
    from jinja2 import Environment, FileSystemLoader
except Exception:
    Environment = None
    FileSystemLoader = None

try:
    from weasyprint import HTML
except Exception:
    HTML = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None

try:
    import qrcode
except Exception:
    qrcode = None

from io import BytesIO
from datetime import datetime
from typing import Dict, Any, Optional
import os


class TemplateEngine:
    """Движок шаблонизации документов"""

    def __init__(self, templates_dir: str = "templates"):
        """Инициализация Jinja2"""
        self.env = Environment(loader=FileSystemLoader(templates_dir))
        self.templates_dir = templates_dir

    def render_template(self, template_name: str, quest_data: Dict[str, Any]) -> str:
        """Рендер HTML шаблона"""
        template = self.env.get_template(template_name)

        # Добавляем текущую дату и QR-код
        context = {
            'quest': quest_data,
            'current_date': datetime.now().strftime("%d.%m.%Y %H:%M"),
            'qr_code_data': self._generate_qr_code(quest_data.get('id', 0))
        }

        return template.render(**context)

    def _generate_qr_code(self, quest_id: int) -> str:
        """Генерация QR-кода с URL квеста"""
        url = f"https://quest-master.local/quest/{quest_id}"

        qr = qrcode.QRCode(version=1, box_size=10, border=2)
        qr.add_data(url)
        qr.make(fit=True)

        img = qr.make_image(fill_color="black", back_color="white")

        # Конвертируем в base64
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        buffer.seek(0)

        import base64
        img_str = base64.b64encode(buffer.getvalue()).decode()

        return f"data:image/png;base64,{img_str}"

    def export_to_pdf(self, template_name: str, quest_data: Dict[str, Any],
                      output_path: Optional[str] = None) -> str:
        """Экспорт в PDF через WeasyPrint"""
        html_content = self.render_template(template_name, quest_data)

        if output_path is None:
            os.makedirs("parchments", exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            quest_id = quest_data.get('id', 'unknown')
            output_path = f"parchments/quest_{quest_id}_{timestamp}.pdf"

        HTML(string=html_content).write_pdf(output_path)
        return output_path

    def export_to_docx(self, quest_data: Dict[str, Any],
                       output_path: Optional[str] = None) -> str:
        """Экспорт в DOCX через python-docx"""
        doc = Document()

        # Заголовок
        title = doc.add_heading(f"Квест: {quest_data.get('title', 'Без названия')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Информация о квесте
        doc.add_paragraph(f"ID квеста: {quest_data.get('id', 'N/A')}")
        doc.add_paragraph(f"Сложность: {quest_data.get('difficulty', 'N/A')}")
        doc.add_paragraph(f"Награда: {quest_data.get('reward', 0)} золотых")
        doc.add_paragraph(f"Дедлайн: {quest_data.get('deadline', 'N/A')}")

        # Описание
        doc.add_heading("Описание квеста:", level=1)
        desc_para = doc.add_paragraph(quest_data.get('description', 'Нет описания'))

        # Печать гильдии
        doc.add_page_break()
        seal = doc.add_paragraph()
        seal.alignment = WD_ALIGN_PARAGRAPH.CENTER
        seal_run = seal.add_run("🏰 ПЕЧАТЬ ГИЛЬДИИ ПРИКЛЮЧЕНЦЕВ 🏰")
        seal_run.bold = True
        seal_run.font.size = Pt(16)

        # Дата создания
        date_para = doc.add_paragraph(f"\nДата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if output_path is None:
            os.makedirs("parchments", exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            quest_id = quest_data.get('id', 'unknown')
            output_path = f"parchments/quest_{quest_id}_{timestamp}.docx"

        doc.save(output_path)
        return output_path


class BatchExporter:
    """Батчевый экспорт для босс-файта"""

    @staticmethod
    def generate_100_quests(db) -> float:
        """
        Генерация 100 квестов для теста производительности
        Returns: время выполнения в секундах
        """
        import time

        start_time = time.time()

        for i in range(100):
            title = f"Тестовый квест #{i+1}"
            difficulty = ["Легкий", "Средний", "Сложный", "Эпический"][i % 4]
            reward = (i + 1) * 100
            description = f"Описание тестового квеста номер {i+1}. " * 10  # 50+ слов
            deadline = "2025-12-31 23:59:59"

            db.create_quest(title, difficulty, reward, description, deadline)

        elapsed = time.time() - start_time
        return elapsed